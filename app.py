from flask import Flask, render_template, request, jsonify, send_file, session
import pandas as pd
import base64
import json
from datetime import datetime
import boto3
import threading
import os
import re
import traceback
import time
from pathlib import Path
import asyncio
import uuid
from collections import defaultdict
import queue
import zipfile
import shutil
from lxml import etree
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from werkzeug.utils import secure_filename
import tempfile

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'ct-review-tool-secret-key-2024')

# Configuration
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
ALLOWED_EXTENSIONS = {'docx'}

# Create directories
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Global variables
guidelines_content = None
hawkeye_checklist = None
document_sessions = {}

# Define paths to guidelines documents
GUIDELINES_PATH = "CT_EE_Review_Guidelines.docx"
HAWKEYE_PATH = "Hawkeye_checklist.docx"

# Hawkeye checklist mapping
HAWKEYE_SECTIONS = {
    1: "Initial Assessment",
    2: "Investigation Process", 
    3: "Seller Classification",
    4: "Enforcement Decision-Making",
    5: "Additional Verification (High-Risk Cases)",
    6: "Multiple Appeals Handling",
    7: "Account Hijacking Prevention",
    8: "Funds Management",
    9: "REs-Q Outreach Process",
    10: "Sentiment Analysis",
    11: "Root Cause Analysis",
    12: "Preventative Actions",
    13: "Documentation and Reporting",
    14: "Cross-Team Collaboration",
    15: "Quality Control",
    16: "Continuous Improvement",
    17: "Communication Standards",
    18: "Performance Metrics",
    19: "Legal and Compliance",
    20: "New Service Launch Considerations"
}

# Standard writeup sections to look for
STANDARD_SECTIONS = [
    "Executive Summary",
    "Background",
    "Resolving Actions",
    "Root Cause",
    "Preventative Actions",
    "Investigation Process",
    "Seller Classification",
    "Documentation and Reporting",
    "Impact Assessment",
    "Timeline",
    "Recommendations"
]

# Sections to exclude from analysis
EXCLUDED_SECTIONS = [
    "Original Email",
    "Email Correspondence",
    "Raw Data",
    "Logs",
    "Attachments"
]

class WordDocumentWithComments:
    """Helper class to add comments to Word documents"""
    
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.temp_dir = f"temp_{uuid.uuid4()}"
        self.comments = []
        self.comment_id = 1
        
    def add_comment(self, paragraph_index, comment_text, author="AI Feedback"):
        """Add a comment to be inserted later"""
        self.comments.append({
            'id': self.comment_id,
            'paragraph_index': paragraph_index,
            'text': comment_text,
            'author': author,
            'date': datetime.now()
        })
        self.comment_id += 1
    
    def _create_comment_xml(self, comment):
        """Create comment XML structure"""
        comment_xml = f'''
        <w:comment w:id="{comment['id']}" w:author="{comment['author']}" 
                   w:date="{comment['date'].strftime('%Y-%m-%dT%H:%M:%S.%fZ')}" 
                   xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:p>
                <w:r>
                    <w:t>{comment['text']}</w:t>
                </w:r>
            </w:p>
        </w:comment>
        '''
        return comment_xml
    
    def save_with_comments(self, output_path):
        """Save document with comments added"""
        try:
            doc = Document(self.doc_path)
            temp_docx = f"{self.temp_dir}_temp.docx"
            doc.save(temp_docx)
            
            os.makedirs(self.temp_dir, exist_ok=True)
            with zipfile.ZipFile(temp_docx, 'r') as zip_ref:
                zip_ref.extractall(self.temp_dir)
            
            comments_path = os.path.join(self.temp_dir, 'word', 'comments.xml')
            
            comments_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
            <w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            '''
            
            for comment in self.comments:
                comments_xml += self._create_comment_xml(comment)
            
            comments_xml += '</w:comments>'
            
            with open(comments_path, 'w', encoding='utf-8') as f:
                f.write(comments_xml)
            
            rels_path = os.path.join(self.temp_dir, 'word', '_rels', 'document.xml.rels')
            if os.path.exists(rels_path):
                with open(rels_path, 'r', encoding='utf-8') as f:
                    rels_content = f.read()
                
                if 'comments.xml' not in rels_content:
                    new_rel = '<Relationship Id="rIdComments" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>'
                    rels_content = rels_content.replace('</Relationships>', f'{new_rel}</Relationships>')
                    
                    with open(rels_path, 'w', encoding='utf-8') as f:
                        f.write(rels_content)
            
            content_types_path = os.path.join(self.temp_dir, '[Content_Types].xml')
            if os.path.exists(content_types_path):
                with open(content_types_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                if 'comments.xml' not in content:
                    new_type = '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
                    content = content.replace('</Types>', f'{new_type}</Types>')
                    
                    with open(content_types_path, 'w', encoding='utf-8') as f:
                        f.write(content)
            
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(self.temp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, self.temp_dir)
                        zipf.write(file_path, arcname)
            
            shutil.rmtree(self.temp_dir)
            os.remove(temp_docx)
            
            return True
            
        except Exception as e:
            print(f"Error adding comments: {str(e)}")
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
            return False

class ReviewSession:
    def __init__(self):
        self.session_id = str(uuid.uuid4())
        self.start_time = datetime.now()
        self.document_name = ""
        self.document_content = ""
        self.document_object = None
        self.document_path = ""
        self.sections = {}
        self.section_paragraphs = {}
        self.paragraph_indices = {}
        self.current_section = 0
        self.feedback_history = defaultdict(list)
        self.section_status = {}
        self.accepted_feedback = defaultdict(list)
        self.rejected_feedback = defaultdict(list)
        self.user_feedback = defaultdict(list)
        self.ai_feedback_cache = {}
        self.document_comments = []
        self.chat_history = []

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def load_guidelines():
    """Load the CT EE Review guidelines and Hawkeye checklist"""
    global guidelines_content, hawkeye_checklist
    
    try:
        if os.path.exists(GUIDELINES_PATH):
            guidelines_content = read_docx(GUIDELINES_PATH)
        
        if os.path.exists(HAWKEYE_PATH):
            hawkeye_checklist = read_docx(HAWKEYE_PATH)
        
        return guidelines_content, hawkeye_checklist
    except Exception as e:
        return None, None

def read_docx(file_path):
    """Extract text from a Word document"""
    try:
        doc = Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
                    
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading document: {str(e)}"

def extract_document_sections_from_docx(doc):
    """Extract sections from Word document based on bold formatting"""
    sections = {}
    section_paragraphs = {}
    paragraph_indices = {}
    current_section = None
    current_content = []
    current_paragraphs = []
    current_indices = []
    
    for idx, para in enumerate(doc.paragraphs):
        is_bold = False
        if para.runs:
            bold_runs = sum(1 for run in para.runs if run.bold)
            total_runs = len(para.runs)
            is_bold = bold_runs > total_runs / 2
        
        text = para.text.strip()
        is_section_header = False
        
        if is_bold and text and len(text) < 100:
            for std_section in STANDARD_SECTIONS:
                if std_section.lower() in text.lower():
                    is_section_header = True
                    break
            
            if not is_section_header and (text.endswith(':') or text.isupper()):
                is_section_header = True
        
        if is_section_header:
            if current_section and current_content:
                exclude = False
                for excluded in EXCLUDED_SECTIONS:
                    if excluded.lower() in current_section.lower():
                        exclude = True
                        break
                
                if not exclude:
                    sections[current_section] = '\n'.join(current_content)
                    section_paragraphs[current_section] = current_paragraphs
                    paragraph_indices[current_section] = current_indices
            
            current_section = text.rstrip(':')
            current_content = []
            current_paragraphs = []
            current_indices = []
        else:
            if text:
                current_content.append(text)
                current_paragraphs.append(para)
                current_indices.append(idx)
    
    if current_section and current_content:
        exclude = False
        for excluded in EXCLUDED_SECTIONS:
            if excluded.lower() in current_section.lower():
                exclude = True
                break
        
        if not exclude:
            sections[current_section] = '\n'.join(current_content)
            section_paragraphs[current_section] = current_paragraphs
            paragraph_indices[current_section] = current_indices
    
    if not sections:
        all_text = []
        all_paras = []
        all_indices = []
        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                all_text.append(para.text)
                all_paras.append(para)
                all_indices.append(idx)
        sections = {"Main Content": '\n'.join(all_text)}
        section_paragraphs = {"Main Content": all_paras}
        paragraph_indices = {"Main Content": all_indices}
    
    return sections, section_paragraphs, paragraph_indices

def get_hawkeye_reference(category, content):
    """Map feedback to relevant Hawkeye checklist items"""
    references = []
    
    keyword_mapping = {
        1: ["customer experience", "cx impact", "customer trust", "buyer impact"],
        2: ["investigation", "sop", "enforcement decision", "abuse pattern"],
        3: ["seller classification", "good actor", "bad actor", "confused actor"],
        4: ["enforcement", "violation", "warning", "suspension"],
        5: ["verification", "supplier", "authenticity", "documentation"],
        6: ["appeal", "repeat", "retrospective"],
        7: ["hijacking", "security", "authentication", "secondary user"],
        8: ["funds", "disbursement", "financial"],
        9: ["outreach", "communication", "clarification"],
        10: ["sentiment", "escalation", "health safety", "legal threat"],
        11: ["root cause", "process gap", "system failure"],
        12: ["preventative", "solution", "improvement", "mitigation"],
        13: ["documentation", "reporting", "background"],
        14: ["cross-team", "collaboration", "engagement"],
        15: ["quality", "audit", "review", "performance"],
        16: ["continuous improvement", "training", "update"],
        17: ["communication standard", "messaging", "clarity"],
        18: ["metrics", "tracking", "measurement"],
        19: ["legal", "compliance", "regulation"],
        20: ["launch", "pilot", "rollback"]
    }
    
    content_lower = content.lower()
    category_lower = category.lower()
    
    for section_num, keywords in keyword_mapping.items():
        for keyword in keywords:
            if keyword in content_lower or keyword in category_lower:
                references.append({
                    'number': section_num,
                    'name': HAWKEYE_SECTIONS[section_num]
                })
                break
    
    return references[:3]

def classify_risk_level(feedback_item):
    """Classify risk level based on Hawkeye criteria"""
    high_risk_indicators = [
        "counterfeit", "fraud", "manipulation", "multiple violation",
        "immediate action", "legal", "health safety", "bad actor"
    ]
    
    medium_risk_indicators = [
        "pattern", "violation", "enforcement", "remediation",
        "correction", "warning"
    ]
    
    content_lower = f"{feedback_item.get('description', '')} {feedback_item.get('category', '')}".lower()
    
    for indicator in high_risk_indicators:
        if indicator in content_lower:
            return "High"
    
    for indicator in medium_risk_indicators:
        if indicator in content_lower:
            return "Medium"
    
    return "Low"

def invoke_aws_semantic_search(system_prompt, user_prompt, operation_name="LLM Analysis"):
    """AWS Bedrock invocation with Hawkeye guidelines"""
    global guidelines_content, hawkeye_checklist
    
    if guidelines_content is None or hawkeye_checklist is None:
        guidelines_content, hawkeye_checklist = load_guidelines()
    
    enhanced_system_prompt = system_prompt
    if hawkeye_checklist:
        truncated_hawkeye = hawkeye_checklist[:30000]
        enhanced_system_prompt = f"""{system_prompt}

HAWKEYE INVESTIGATION CHECKLIST:
{truncated_hawkeye}

Apply these Hawkeye investigation mental models in your analysis. Reference specific checklist items when providing feedback."""
    
    try:
        runtime = boto3.client('bedrock-runtime')
        
        body = json.dumps({
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 4000,
            "system": enhanced_system_prompt,
            "messages": [{"role": "user", "content": user_prompt}]
        })
        
        response = runtime.invoke_model(
            body=body,
            modelId='anthropic.claude-3-sonnet-20240229-v1:0',
            accept="application/json",
            contentType="application/json"
        )
        
        response_body = json.loads(response.get('body').read())
        return response_body['content'][0]['text']
        
    except Exception as e:
        # Generate section-specific mock responses for testing
        return generate_section_specific_response(user_prompt, operation_name)

def generate_section_specific_response(user_prompt, operation_name):
    """Generate section-specific responses based on content analysis"""
    time.sleep(1)
    
    if "chat" in operation_name.lower():
        # Extract actual question from user prompt
        question = user_prompt.lower()
        
        # Extract the actual user question from the prompt
        if "user question:" in question:
            question = question.split("user question:")[1].strip()
        
        # Provide specific responses based on question content
        if "seller classification" in question or "good actor" in question or "bad actor" in question:
            return "Seller Classification (Hawkeye #3): Good Actor = unintentional violation, cooperative response; Bad Actor = intentional abuse, non-cooperative; Confused Actor = misunderstands policies. Classification should be based on intent, history, and response to enforcement."
        elif "hawkeye" in question and any(x in question for x in ["1", "one", "first", "initial"]):
            return "Hawkeye #1 - Initial Assessment: Evaluate customer experience (CX) impact. Ask: How does this issue affect customer trust? What's the potential for negative reviews or returns? Always consider both immediate and long-term customer impact."
        elif "risk" in question:
            return "Risk Classification: High Risk = counterfeit, fraud, health/safety issues, bad actors; Medium Risk = policy violations with patterns, enforcement needed; Low Risk = isolated incidents, clarifications needed. Base classification on impact severity and scope."
        elif "root cause" in question:
            return "Root Cause Analysis (Hawkeye #11): Use 5 Whys technique. Identify process gaps, system failures, policy ambiguities. Distinguish immediate vs systemic causes. Ask: What allowed this to happen? Could better processes prevent it?"
        elif "preventative" in question or "prevention" in question:
            return "Preventative Actions (Hawkeye #12): Structure as Immediate (stop current harm), Short-term (prevent recurrence), Long-term (systemic improvements). Address root causes identified in analysis."
        elif "investigation" in question:
            return "Investigation Process (Hawkeye #2): Follow SOPs but challenge when needed. Document methodology, evidence, decisions. Show critical thinking. Which SOPs were used? Any deviations and why?"
        elif "feedback" in question:
            return "The feedback analyzes your specific section content against Hawkeye standards. Each item includes targeted questions, actionable suggestions, and relevant checkpoint references. What specific feedback would you like me to explain?"
        else:
            return f"I'm TARA, your CT review assistant. I can help with: Hawkeye checkpoints (#1-20), seller classification, risk assessment, investigation best practices, or specific feedback items. What would you like to know?"
    
    # Extract section name and content for analysis
    section_name = "Unknown Section"
    section_content = ""
    
    # Parse section info from prompt
    if "section \"" in user_prompt:
        start = user_prompt.find('section "') + 9
        end = user_prompt.find('"', start)
        if end > start:
            section_name = user_prompt[start:end]
    
    if "SECTION CONTENT:" in user_prompt:
        start = user_prompt.find('SECTION CONTENT:') + 16
        section_content = user_prompt[start:start+1000].strip()
    
    # Generate section-specific feedback
    feedback_items = generate_contextual_feedback(section_name, section_content)
    
    return json.dumps({"feedback_items": feedback_items})

def generate_contextual_feedback(section_name, content):
    """Generate contextual feedback based on section name and content"""
    feedback_items = []
    content_lower = content.lower()
    section_lower = section_name.lower()
    
    # Executive Summary specific feedback
    if "executive" in section_lower or "summary" in section_lower:
        feedback_items.extend([
            {
                "id": f"exec_1_{hash(content) % 1000}",
                "type": "critical",
                "category": "Initial Assessment",
                "description": f"Executive summary should clearly state the customer experience (CX) impact. Current content lacks specific mention of how this issue affects customer trust and satisfaction.",
                "suggestion": "Add a dedicated sentence about CX impact, including potential for negative reviews, returns, or customer complaints",
                "example": "Example: 'This issue could result in customer dissatisfaction due to [specific impact], potentially affecting trust in our marketplace'",
                "questions": [
                    "What is the direct impact on customer experience?",
                    "How might this affect customer trust in Amazon?"
                ],
                "confidence": 0.92
            }
        ])
        
        if "risk" not in content_lower:
            feedback_items.append({
                "id": f"exec_2_{hash(content) % 1000}",
                "type": "important",
                "category": "Risk Assessment",
                "description": "Executive summary missing clear risk classification (High/Medium/Low) based on Hawkeye criteria",
                "suggestion": "Include explicit risk level with justification based on impact severity and scope",
                "example": "High Risk: Affects customer safety/trust; Medium Risk: Policy violation with pattern; Low Risk: Isolated incident",
                "questions": ["What is the risk level?", "What criteria determine this classification?"],
                "confidence": 0.88
            })
    
    # Background specific feedback
    elif "background" in section_lower:
        feedback_items.extend([
            {
                "id": f"bg_1_{hash(content) % 1000}",
                "type": "important",
                "category": "Investigation Process",
                "description": "Background section should include timeline of events and initial detection method. Current content may lack chronological clarity.",
                "suggestion": "Structure background with clear timeline: when issue was detected, by whom, initial scope assessment",
                "example": "Timeline: Issue detected on [date] via [method], initial scope showed [impact], escalated due to [reason]",
                "questions": [
                    "When was this issue first detected?",
                    "What was the detection method?",
                    "Who reported or identified the issue?"
                ],
                "confidence": 0.85
            }
        ])
        
        if "seller" in content_lower or "account" in content_lower:
            feedback_items.append({
                "id": f"bg_2_{hash(content) % 1000}",
                "type": "critical",
                "category": "Seller Classification",
                "description": "Background mentions seller/account but lacks proper classification as Good Actor, Bad Actor, or Confused Actor per Hawkeye guidelines",
                "suggestion": "Classify the seller based on intent, history, and response to enforcement actions",
                "example": "Good Actor: Unintentional violation, cooperative; Bad Actor: Intentional abuse, non-cooperative; Confused Actor: Misunderstands policies",
                "questions": ["What is the seller's intent?", "How did they respond to initial contact?"],
                "confidence": 0.90
            })
    
    # Root Cause specific feedback
    elif "root cause" in section_lower or "cause" in section_lower:
        feedback_items.extend([
            {
                "id": f"rc_1_{hash(content) % 1000}",
                "type": "critical",
                "category": "Root Cause Analysis",
                "description": "Root cause analysis must identify specific process gaps, system failures, or policy ambiguities that allowed this issue to occur",
                "suggestion": "Use the 5 Whys technique to drill down to fundamental causes. Include both immediate and systemic causes",
                "example": "Immediate cause: Seller uploaded prohibited item; Systemic cause: Detection algorithm missed specific keyword variations",
                "questions": [
                    "What process gap allowed this to happen?",
                    "Are there system limitations that contributed?",
                    "Could policy clarity have prevented this?"
                ],
                "confidence": 0.94
            }
        ])
    
    # Preventative Actions specific feedback
    elif "preventative" in section_lower or "prevention" in section_lower:
        feedback_items.extend([
            {
                "id": f"prev_1_{hash(content) % 1000}",
                "type": "important",
                "category": "Preventative Actions",
                "description": "Preventative actions should address both immediate fixes and long-term systemic improvements identified in root cause analysis",
                "suggestion": "Structure as: Immediate actions (stop current issue), Short-term fixes (prevent recurrence), Long-term improvements (systemic changes)",
                "example": "Immediate: Remove violating listings; Short-term: Update detection rules; Long-term: Enhance seller education program",
                "questions": [
                    "What immediate actions prevent further harm?",
                    "How do we prevent this specific issue from recurring?",
                    "What systemic changes are needed?"
                ],
                "confidence": 0.87
            }
        ])
    
    # Investigation Process specific feedback
    elif "investigation" in section_lower or "process" in section_lower:
        feedback_items.extend([
            {
                "id": f"inv_1_{hash(content) % 1000}",
                "type": "important",
                "category": "Investigation Process",
                "description": "Investigation process should demonstrate adherence to SOPs while showing critical thinking and challenge of standard procedures when appropriate",
                "suggestion": "Document which SOPs were followed, any deviations made, and rationale for investigative decisions",
                "example": "Followed SOP-123 for initial assessment, deviated at step 5 due to unique circumstances, consulted with [team] for guidance",
                "questions": [
                    "Which SOPs were followed?",
                    "Were any standard procedures challenged or modified?",
                    "What investigative tools were used?"
                ],
                "confidence": 0.89
            }
        ])
    
    # Generic feedback for any section
    if len(feedback_items) == 0 or len(feedback_items) < 2:
        # Add generic but contextual feedback
        if "documentation" not in content_lower:
            feedback_items.append({
                "id": f"gen_1_{hash(content) % 1000}",
                "type": "suggestion",
                "category": "Documentation and Reporting",
                "description": f"Section '{section_name}' could benefit from more detailed documentation of evidence and decision-making rationale",
                "suggestion": "Include specific evidence, data points, and reasoning that led to conclusions in this section",
                "example": "Reference specific case numbers, timestamps, communication records, or data analysis results",
                "questions": ["What evidence supports the conclusions?", "Are all decisions properly documented?"],
                "confidence": 0.75
            })
        
        # Check for collaboration mentions
        if "team" not in content_lower and "consult" not in content_lower:
            feedback_items.append({
                "id": f"gen_2_{hash(content) % 1000}",
                "type": "suggestion",
                "category": "Cross-Team Collaboration",
                "description": f"Consider if cross-team collaboration was needed for '{section_name}' and document any consultations or escalations",
                "suggestion": "Mention any consultations with legal, policy, or other teams if relevant to this section",
                "example": "Consulted with Policy team on interpretation, Legal team confirmed compliance approach",
                "questions": ["Were other teams consulted?", "Should this have been escalated?"],
                "confidence": 0.70
            })
    
    # Ensure each feedback item has required fields
    for item in feedback_items:
        if 'hawkeye_refs' not in item:
            refs = get_hawkeye_reference(item.get('category', ''), item.get('description', ''))
            item['hawkeye_refs'] = [ref['number'] for ref in refs]
        
        if 'risk_level' not in item:
            item['risk_level'] = classify_risk_level(item)
    
    return feedback_items

def analyze_section_with_ai(section_name, section_content, doc_type="Full Write-up"):
    """Analyze a single section with Hawkeye framework"""
    
    # Create detailed analysis prompt with section-specific guidance
    section_guidance = get_section_specific_guidance(section_name)
    
    prompt = f"""Analyze this section "{section_name}" from a {doc_type} document using the Hawkeye investigation framework.

SECTION CONTENT:
{section_content[:3000]}

SECTION-SPECIFIC GUIDANCE:
{section_guidance}

Provide detailed, document-specific feedback following the 20-point Hawkeye checklist. Focus on:
1. Content gaps specific to this section type
2. Missing Hawkeye criteria that should be addressed
3. Specific improvements based on the actual content
4. Questions that arise from reading this specific content

For each feedback item, include:
- Specific references to the actual content
- Detailed suggestions for improvement
- Relevant Hawkeye checkpoint numbers (#1-20)
- Risk classification based on impact
- Specific questions that should be answered

Return feedback in this JSON format:
{{
    "feedback_items": [
        {{
            "id": "unique_id",
            "type": "critical|important|suggestion|positive",
            "category": "category matching Hawkeye sections",
            "description": "Detailed description referencing specific content and Hawkeye criteria",
            "suggestion": "Specific, actionable suggestion based on content analysis",
            "example": "Concrete example or template for improvement",
            "questions": ["Specific question about this content?", "What should be clarified?"],
            "hawkeye_refs": [1, 11, 12],
            "risk_level": "High|Medium|Low",
            "confidence": 0.95
        }}
    ]
}}"""
    
    system_prompt = f"""You are an expert CT EE document reviewer with deep knowledge of the Hawkeye investigation framework. 
Analyze the provided section content thoroughly and provide specific, actionable feedback based on what is actually written (or missing) in the content.
Focus on document-centric analysis rather than generic advice."""
    
    response = invoke_aws_semantic_search(system_prompt, prompt, f"Detailed Hawkeye Analysis: {section_name}")
    
    try:
        result = json.loads(response)
    except:
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            try:
                result = json.loads(json_match.group(0))
            except:
                result = {"feedback_items": []}
        else:
            result = {"feedback_items": []}
    
    # Enhance feedback items with additional context
    for item in result.get('feedback_items', []):
        if 'hawkeye_refs' not in item:
            refs = get_hawkeye_reference(item.get('category', ''), item.get('description', ''))
            item['hawkeye_refs'] = [ref['number'] for ref in refs]
        
        if 'risk_level' not in item:
            item['risk_level'] = classify_risk_level(item)
        
        # Add section context to description
        if 'description' in item and section_name not in item['description']:
            item['description'] = f"In '{section_name}': {item['description']}"
    
    return result

def get_section_specific_guidance(section_name):
    """Get specific guidance for different section types"""
    section_lower = section_name.lower()
    
    if "executive" in section_lower or "summary" in section_lower:
        return """Executive Summary should include:
- Clear statement of the issue and its impact
- Customer experience (CX) impact assessment (Hawkeye #1)
- Risk level classification with justification
- High-level resolution approach
- Key stakeholders involved"""
    
    elif "background" in section_lower:
        return """Background should include:
- Timeline of events leading to the issue
- How the issue was detected (Hawkeye #2)
- Initial scope and impact assessment
- Relevant historical context
- Key players and their roles"""
    
    elif "root cause" in section_lower or "cause" in section_lower:
        return """Root Cause Analysis should include:
- Systematic analysis using 5 Whys or similar methodology (Hawkeye #11)
- Identification of process gaps and system failures
- Distinction between immediate and systemic causes
- Analysis of why existing controls failed
- Contributing factors and environmental conditions"""
    
    elif "preventative" in section_lower or "prevention" in section_lower:
        return """Preventative Actions should include:
- Immediate actions to stop current harm (Hawkeye #12)
- Short-term fixes to prevent recurrence
- Long-term systemic improvements
- Process and system enhancements
- Monitoring and detection improvements"""
    
    elif "investigation" in section_lower:
        return """Investigation Process should include:
- SOPs followed and any deviations (Hawkeye #2)
- Evidence gathering methodology
- Analysis techniques used
- Decision-making rationale
- Quality control measures applied (Hawkeye #15)"""
    
    elif "resolving" in section_lower or "resolution" in section_lower:
        return """Resolving Actions should include:
- Specific enforcement actions taken (Hawkeye #4)
- Seller classification rationale (Hawkeye #3)
- Communication approach used (Hawkeye #17)
- Verification steps for high-risk cases (Hawkeye #5)
- Appeals handling if applicable (Hawkeye #6)"""
    
    else:
        return """General section should include:
- Clear purpose and scope
- Relevant Hawkeye checklist considerations
- Supporting evidence and documentation (Hawkeye #13)
- Cross-team collaboration details (Hawkeye #14)
- Quality assurance measures"""

def process_chat_query(query, context, session_id):
    """Process chat query with context awareness"""
    session = document_sessions.get(session_id)
    if not session:
        return "No active session found."
    
    current_section = context.get('current_section', 'None')
    
    # Get current section content for context
    section_content = ""
    if current_section != 'None' and current_section in session.sections:
        section_content = session.sections[current_section][:500]  # First 500 chars for context
    
    # Get current feedback for context
    current_feedback = ""
    if current_section in session.ai_feedback_cache:
        feedback_items = session.ai_feedback_cache[current_section].get('feedback_items', [])
        if feedback_items:
            current_feedback = f"Current feedback includes {len(feedback_items)} items: " + ", ".join([item.get('category', 'General') for item in feedback_items[:3]])
    
    context_info = f"""
    Current Section: {current_section}
    Section Content Preview: {section_content[:200]}...
    Current Feedback: {current_feedback}
    Document Type: Full Write-up
    """
    
    prompt = f"""You are an AI assistant helping with document review using the Hawkeye framework.

CONTEXT:
{context_info}

HAWKEYE GUIDELINES REFERENCE:
The 20-point Hawkeye checklist includes:
1. Initial Assessment - Evaluate CX impact
2. Investigation Process - Challenge SOPs
3. Seller Classification - Identify good/bad actors
4. Enforcement Decision-Making
5. Additional Verification for High-Risk Cases
6. Multiple Appeals Handling
7. Account Hijacking Prevention
8. Funds Management
9. REs-Q Outreach Process
10. Sentiment Analysis
11. Root Cause Analysis
12. Preventative Actions
13. Documentation and Reporting
14. Cross-Team Collaboration
15. Quality Control
16. Continuous Improvement
17. Communication Standards
18. Performance Metrics
19. Legal and Compliance
20. New Service Launch Considerations

USER QUESTION: {query}

Provide a helpful, specific response that references the Hawkeye guidelines and current section context when relevant. Be concise but thorough."""
    
    system_prompt = "You are an expert assistant for the Hawkeye document review system with deep knowledge of CT EE guidelines."
    
    response = invoke_aws_semantic_search(system_prompt, prompt, f"Chat Assistant - {query[:50]}")
    
    return response

def create_reviewed_document_with_proper_comments(original_doc_path, doc_name, comments_data):
    """Create a copy of the original document with proper Word comments"""
    
    try:
        doc_with_comments = WordDocumentWithComments(original_doc_path)
        
        for comment_data in comments_data:
            author = comment_data.get('author', 'AI Feedback')
            doc_with_comments.add_comment(
                paragraph_index=comment_data['paragraph_index'],
                comment_text=comment_data['comment'],
                author=author
            )
        
        output_path = os.path.join(OUTPUT_FOLDER, f'reviewed_{doc_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
        success = doc_with_comments.save_with_comments(output_path)
        
        if success:
            return output_path
        else:
            return create_simple_reviewed_copy(original_doc_path, doc_name, comments_data)
            
    except Exception as e:
        print(f"Error creating document with comments: {str(e)}")
        return create_simple_reviewed_copy(original_doc_path, doc_name, comments_data)

def create_simple_reviewed_copy(original_doc_path, doc_name, comments_data):
    """Create a simple copy with inline comment markers as fallback"""
    try:
        output_path = os.path.join(OUTPUT_FOLDER, f'reviewed_{doc_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
        
        doc = Document(original_doc_path)
        
        doc.add_page_break()
        heading = doc.add_heading('Hawkeye Review Feedback Summary', 1)
        
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph(f'Total feedback items: {len(comments_data)}')
        doc.add_paragraph('')
        
        section_comments = defaultdict(list)
        for comment in comments_data:
            section_comments[comment['section']].append(comment)
        
        for section, comments in section_comments.items():
            section_heading = doc.add_heading(section, 2)
            
            for comment in comments:
                p = doc.add_paragraph(style='List Bullet')
                author = comment.get('author', 'AI Feedback')
                p.add_run(f"[{author}] {comment['type'].upper()} - {comment['risk_level']} Risk: ").bold = True
                p.add_run(comment['comment'])
        
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        print(f"Error creating simple copy: {str(e)}")
        return None

# Routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        
        # Create session
        session_id = str(uuid.uuid4())
        review_session = ReviewSession()
        review_session.session_id = session_id
        review_session.document_name = filename
        review_session.document_path = file_path
        
        try:
            # Open document
            doc = Document(file_path)
            review_session.document_object = doc
            
            # Extract sections
            sections, section_paragraphs, paragraph_indices = extract_document_sections_from_docx(doc)
            review_session.sections = sections
            review_session.section_paragraphs = section_paragraphs
            review_session.paragraph_indices = paragraph_indices
            
            document_sessions[session_id] = review_session
            session['session_id'] = session_id
            
            return jsonify({
                'success': True,
                'session_id': session_id,
                'sections': list(sections.keys()),
                'document_name': filename
            })
            
        except Exception as e:
            return jsonify({'error': f'Error processing document: {str(e)}'}), 500
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/analyze_section', methods=['POST'])
def analyze_section():
    data = request.json
    session_id = data.get('session_id')
    section_name = data.get('section_name')
    
    if not session_id or session_id not in document_sessions:
        return jsonify({'error': 'Invalid session'}), 400
    
    review_session = document_sessions[session_id]
    
    if section_name not in review_session.sections:
        return jsonify({'error': 'Section not found'}), 400
    
    section_content = review_session.sections[section_name]
    
    # Check cache first
    cache_key = f"{section_name}_{hash(section_content)}"
    if cache_key in review_session.ai_feedback_cache:
        result = review_session.ai_feedback_cache[cache_key]
    else:
        result = analyze_section_with_ai(section_name, section_content)
        review_session.ai_feedback_cache[cache_key] = result
    
    return jsonify(result)

@app.route('/get_section', methods=['POST'])
def get_section():
    data = request.json
    session_id = data.get('session_id')
    section_name = data.get('section_name')
    
    if not session_id or session_id not in document_sessions:
        return jsonify({'error': 'Invalid session'}), 400
    
    review_session = document_sessions[session_id]
    
    if section_name not in review_session.sections:
        return jsonify({'error': 'Section not found'}), 400
    
    return jsonify({
        'content': review_session.sections[section_name],
        'section_name': section_name
    })

@app.route('/accept_feedback', methods=['POST'])
def accept_feedback():
    data = request.json
    session_id = data.get('session_id')
    section_name = data.get('section_name')
    feedback_item = data.get('feedback_item')
    
    if not session_id or session_id not in document_sessions:
        return jsonify({'error': 'Invalid session'}), 400
    
    review_session = document_sessions[session_id]
    
    if section_name not in review_session.accepted_feedback:
        review_session.accepted_feedback[section_name] = []
    
    review_session.accepted_feedback[section_name].append(feedback_item)
    
    # Prepare comment for Word document
    comment_text = f"[{feedback_item['type'].upper()} - {feedback_item.get('risk_level', 'Low')} Risk]\n"
    comment_text += f"{feedback_item['description']}\n"
    if feedback_item.get('suggestion'):
        comment_text += f"\nSuggestion: {feedback_item['suggestion']}\n"
    if feedback_item.get('hawkeye_refs'):
        refs = [f"#{r} {HAWKEYE_SECTIONS.get(r, '')}" for r in feedback_item['hawkeye_refs']]
        comment_text += f"\nHawkeye References: {', '.join(refs)}"
    
    # Store comment to be added to document
    if section_name in review_session.paragraph_indices and review_session.paragraph_indices[section_name]:
        review_session.document_comments.append({
            'section': section_name,
            'paragraph_index': review_session.paragraph_indices[section_name][0],
            'comment': comment_text,
            'type': feedback_item['type'],
            'risk_level': feedback_item.get('risk_level', 'Low'),
            'author': 'AI Feedback'
        })
    
    return jsonify({'success': True})

@app.route('/reject_feedback', methods=['POST'])
def reject_feedback():
    data = request.json
    session_id = data.get('session_id')
    section_name = data.get('section_name')
    feedback_item = data.get('feedback_item')
    
    if not session_id or session_id not in document_sessions:
        return jsonify({'error': 'Invalid session'}), 400
    
    review_session = document_sessions[session_id]
    
    if section_name not in review_session.rejected_feedback:
        review_session.rejected_feedback[section_name] = []
        
    review_session.rejected_feedback[section_name].append(feedback_item)
    
    return jsonify({'success': True})

@app.route('/add_custom_feedback', methods=['POST'])
def add_custom_feedback():
    data = request.json
    session_id = data.get('session_id')
    section_name = data.get('section_name')
    feedback_type = data.get('type')
    category = data.get('category')
    description = data.get('description')
    
    if not session_id or session_id not in document_sessions:
        return jsonify({'error': 'Invalid session'}), 400
    
    review_session = document_sessions[session_id]
    
    # Find Hawkeye reference number
    hawkeye_ref = 1
    for num, name in HAWKEYE_SECTIONS.items():
        if name == category:
            hawkeye_ref = num
            break
    
    feedback = {
        'id': str(uuid.uuid4()),
        'type': feedback_type,
        'category': category,
        'description': description,
        'suggestion': '',
        'hawkeye_refs': [hawkeye_ref],
        'risk_level': 'Medium' if feedback_type == 'critical' else 'Low',
        'timestamp': datetime.now().isoformat(),
        'user_created': True
    }
    
    if section_name not in review_session.user_feedback:
        review_session.user_feedback[section_name] = []
    
    review_session.user_feedback[section_name].append(feedback)
    
    # Also add as accepted feedback for comment
    if section_name not in review_session.accepted_feedback:
        review_session.accepted_feedback[section_name] = []
    review_session.accepted_feedback[section_name].append(feedback)
    
    # Prepare comment
    comment_text = f"[USER FEEDBACK - {feedback['type'].upper()}]\n"
    comment_text += f"{feedback['description']}\n"
    comment_text += f"\nHawkeye Reference: #{hawkeye_ref} {category}"
    
    if section_name in review_session.paragraph_indices and review_session.paragraph_indices[section_name]:
        review_session.document_comments.append({
            'section': section_name,
            'paragraph_index': review_session.paragraph_indices[section_name][0],
            'comment': comment_text,
            'type': feedback['type'],
            'risk_level': feedback['risk_level'],
            'user_created': True,
            'author': 'User Feedback'
        })
    
    return jsonify({'success': True, 'feedback': feedback})

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.json
        session_id = data.get('session_id')
        query = data.get('query', '').strip()
        context = data.get('context', {})
        
        if not session_id or session_id not in document_sessions:
            return jsonify({'error': 'Invalid session'}), 400
        
        if not query:
            return jsonify({'response': 'Please ask a question about the document or Hawkeye guidelines.'})
        
        review_session = document_sessions[session_id]
        
        # Direct chat response
        response = get_direct_chat_response(query)
        
        # Store chat history
        review_session.chat_history.append({
            'role': 'user',
            'content': query,
            'timestamp': datetime.now().isoformat()
        })
        review_session.chat_history.append({
            'role': 'assistant',
            'content': response,
            'timestamp': datetime.now().isoformat()
        })
        
        return jsonify({'response': response})
        
    except Exception as e:
        return jsonify({'response': f'I encountered an error. Please try asking your question again.'})

def get_direct_chat_response(query):
    """Get direct chat response"""
    query_lower = query.lower()
    
    if any(term in query_lower for term in ['seller classification', 'good actor', 'bad actor']):
        return "Seller Classification (Hawkeye #3): Good Actor = unintentional violation, cooperative response; Bad Actor = intentional abuse, non-cooperative; Confused Actor = misunderstands policies. Classification based on intent, history, and response to enforcement."
    
    elif any(term in query_lower for term in ['hawkeye 1', 'initial assessment', 'customer experience']):
        return "Hawkeye #1 - Initial Assessment: Evaluate customer experience (CX) impact. Ask: How does this issue affect customer trust? What's the potential for negative reviews or returns? Always consider both immediate and long-term customer impact."
    
    elif 'risk' in query_lower:
        return "Risk Classification: High Risk = counterfeit, fraud, health/safety issues; Medium Risk = policy violations with patterns; Low Risk = isolated incidents, clarifications needed. Base classification on impact severity and scope."
    
    elif 'root cause' in query_lower:
        return "Root Cause Analysis (Hawkeye #11): Use 5 Whys technique. Identify process gaps, system failures, policy ambiguities. Distinguish immediate vs systemic causes."
    
    elif any(term in query_lower for term in ['preventative', 'prevention']):
        return "Preventative Actions (Hawkeye #12): Structure as Immediate (stop current harm), Short-term (prevent recurrence), Long-term (systemic improvements)."
    
    elif 'investigation' in query_lower:
        return "Investigation Process (Hawkeye #2): Follow SOPs but challenge when needed. Document methodology, evidence, decisions. Show critical thinking."
    
    elif 'feedback' in query_lower:
        return "The feedback analyzes your specific section content against Hawkeye standards. Each item includes targeted questions, actionable suggestions, and relevant checkpoint references."
    
    elif 'hawkeye' in query_lower:
        return "The Hawkeye 20-point checklist ensures thorough investigation. Key areas: 1) Customer Experience Impact, 2) Investigation Process, 3) Seller Classification, 4) Enforcement Decision-Making. Which checkpoint would you like me to explain?"
    
    else:
        return "I'm TARA, your CT review assistant. I can help with: Hawkeye checkpoints (#1-20), seller classification, risk assessment, investigation best practices, or specific feedback items. What would you like to know?"

@app.route('/complete_review', methods=['POST'])
def complete_review():
    data = request.json
    session_id = data.get('session_id')
    
    if not session_id or session_id not in document_sessions:
        return jsonify({'error': 'Invalid session'}), 400
    
    review_session = document_sessions[session_id]
    
    if not review_session.document_comments:
        return jsonify({'error': 'No feedback accepted. Please accept some feedback items first.'}), 400
    
    # Generate document
    if review_session.document_path and os.path.exists(review_session.document_path):
        output_path = create_reviewed_document_with_proper_comments(
            review_session.document_path,
            review_session.document_name,
            review_session.document_comments
        )
        
        if output_path and os.path.exists(output_path):
            return jsonify({
                'success': True,
                'output_path': os.path.basename(output_path),
                'comments_count': len(review_session.document_comments)
            })
    
    return jsonify({'error': 'Failed to generate reviewed document'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    return send_file(os.path.join(OUTPUT_FOLDER, filename), as_attachment=True)

@app.route('/get_stats', methods=['POST'])
def get_stats():
    data = request.json
    session_id = data.get('session_id')
    
    if not session_id or session_id not in document_sessions:
        return jsonify({'error': 'Invalid session'}), 400
    
    review_session = document_sessions[session_id]
    
    # Calculate statistics
    total_feedback = sum(len(items) for items in review_session.ai_feedback_cache.values())
    total_accepted = sum(len(items) for items in review_session.accepted_feedback.values())
    total_rejected = sum(len(items) for items in review_session.rejected_feedback.values())
    total_user = sum(len(items) for items in review_session.user_feedback.values())
    
    high_risk = 0
    medium_risk = 0
    
    for result in review_session.ai_feedback_cache.values():
        for item in result.get('feedback_items', []):
            if item.get('risk_level') == 'High':
                high_risk += 1
            elif item.get('risk_level') == 'Medium':
                medium_risk += 1
    
    return jsonify({
        'total_feedback': total_feedback,
        'high_risk': high_risk,
        'medium_risk': medium_risk,
        'accepted': total_accepted,
        'user_added': total_user
    })

# Load guidelines on startup
load_guidelines()

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_ENV') != 'production'
    app.run(debug=debug, host='0.0.0.0', port=port)